"""
Generate the Livo Pronunciation Coach system architecture doc (DOCX).

Usage:
    pip install python-docx
    python docs/generate_architecture_doc.py

Output:
    docs/System_Architecture.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


# --- palette (ink on paper, not "AI purple") ---
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
ACCENT = RGBColor(0x1F, 0x4E, 0x3D)  # deep green
RULE = "1F4E3D"
LIGHT_BG = "F4F6F4"
BOX_BG = "E8EEE9"
ARROW_BG = "FFFFFF"

LIVE_APP = "https://pronunciation-coach.vercel.app"
LIVE_API = "https://pronunciation-coach-api.onrender.com"
GITHUB = "https://github.com/Tusharkanta407/Pronunciation-Coach"


def set_run_font(run, *, size=10.5, bold=False, italic=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_horizontal_rule(paragraph, color=RULE):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), kwargs[edge].get("val", "single"))
            element.set(qn("w:sz"), kwargs[edge].get("sz", "4"))
            element.set(qn("w:color"), kwargs[edge].get("color", RULE))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)


def fill_box_cell(cell, lines: list[str], *, bg=BOX_BG, title=None, center=True):
    clear_cell(cell)
    shade_cell(cell, bg)
    set_cell_border(
        cell,
        top={"sz": "12", "color": RULE},
        bottom={"sz": "12", "color": RULE},
        left={"sz": "12", "color": RULE},
        right={"sz": "12", "color": RULE},
    )
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT

    if title:
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_run_font(r, size=9, bold=True, color=ACCENT)
        for line in lines:
            p2 = cell.add_paragraph()
            p2.alignment = align
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(1)
            r2 = p2.add_run(line)
            set_run_font(r2, size=8.5, color=INK)
    else:
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for i, line in enumerate(lines):
            if i == 0:
                r = p.add_run(line)
            else:
                p2 = cell.add_paragraph()
                p2.alignment = align
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(1)
                r = p2.add_run(line)
            set_run_font(r, size=8.5, color=INK)


def arrow_cell(cell, text="↓"):
    clear_cell(cell)
    shade_cell(cell, ARROW_BG)
    set_cell_border(
        cell,
        top={"sz": "0", "color": "FFFFFF", "val": "nil"},
        bottom={"sz": "0", "color": "FFFFFF", "val": "nil"},
        left={"sz": "0", "color": "FFFFFF", "val": "nil"},
        right={"sz": "0", "color": "FFFFFF", "val": "nil"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=True, color=ACCENT)


def para(doc, text, *, size=10.5, bold=False, italic=False, color=INK, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level == 1 else 6)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12 if level == 1 else 11, bold=True, color=ACCENT)
    if level == 1:
        add_horizontal_rule(p)
    return p


def bullet(doc, text, *, size=10):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=size, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=INK)
    return p


def link_line(doc, label: str, url: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=9.5, bold=True, color=ACCENT)
    r2 = p.add_run(url)
    set_run_font(r2, size=9.5, color=RGBColor(0x1A, 0x5C, 0x96))
    return p


def add_current_architecture_diagram(doc):
    """Vertical box stack: Browser → Vercel → Render → services."""
    para(doc, "Current system (what ships today)", size=10, bold=True, color=ACCENT, space_after=4)

    # Row 1: Browser
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    fill_box_cell(
        t.cell(0, 0),
        ["Mic / file upload", "Consent → 3 rounds → practice"],
        title="①  BROWSER",
    )

    # arrow
    a = doc.add_table(rows=1, cols=1)
    arrow_cell(a.cell(0, 0), "↓  HTTPS")

    # Row 2: Vercel
    t = doc.add_table(rows=1, cols=1)
    fill_box_cell(
        t.cell(0, 0),
        [
            "Next.js UI  ·  Vercel",
            "Rewrites  /api/*  →  Render backend",
            LIVE_APP,
        ],
        title="②  FRONTEND",
    )

    a = doc.add_table(rows=1, cols=1)
    arrow_cell(a.cell(0, 0), "↓  proxy  /api/analyze , /api/session , …")

    # Row 3: FastAPI
    t = doc.add_table(rows=1, cols=1)
    fill_box_cell(
        t.cell(0, 0),
        [
            "FastAPI  ·  Render (Docker + ffmpeg)",
            "Session store in memory  ·  no audio on disk",
            LIVE_API,
        ],
        title="③  BACKEND",
    )

    a = doc.add_table(rows=1, cols=1)
    arrow_cell(a.cell(0, 0), "↓  calls out for speech + feedback")

    # Row 4: 4 service boxes side by side
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    fill_box_cell(t.cell(0, 0), ["whisper-large-v3", "words + confidence"], title="Groq STT", bg=LIGHT_BG)
    fill_box_cell(t.cell(0, 1), ["phoneme distance", "trap-word tips"], title="g2p-en", bg=LIGHT_BG)
    fill_box_cell(t.cell(0, 2), ["hear target word", "practice drill"], title="edge-tts", bg=LIGHT_BG)
    fill_box_cell(t.cell(0, 3), ["1-line tip only", "optional"], title="OpenRouter", bg=LIGHT_BG)

    para(doc, "", size=4, space_after=2)
    para(
        doc,
        "Read left→right on the bottom row: STT turns audio into words; g2p compares "
        "sounds on text; TTS plays the correct word; LLM is just a short tip.",
        size=9,
        italic=True,
        color=MUTED,
        space_after=6,
    )


def add_next_week_diagram(doc):
    """Box diagram for proper pronunciation pipeline."""
    para(
        doc,
        "If I had one more week — proper pronunciation pipeline",
        size=10,
        bold=True,
        color=ACCENT,
        space_after=4,
    )

    t = doc.add_table(rows=1, cols=1)
    fill_box_cell(
        t.cell(0, 0),
        ["Same as today: record → Vercel → FastAPI"],
        title="KEEP  (Layer 1 — Alignment)",
        bg=LIGHT_BG,
    )

    a = doc.add_table(rows=1, cols=1)
    arrow_cell(a.cell(0, 0), "↓  then, only for flagged / trap / practice words")

    # 3 boxes
    t = doc.add_table(rows=1, cols=3)
    fill_box_cell(
        t.cell(0, 0),
        ["cut audio by", "word timestamps", "start → end"],
        title="A  Clip word",
    )
    fill_box_cell(
        t.cell(0, 1),
        ["Azure Pronunciation", "Assessment  (or", "Speechace)"],
        title="B  Acoustic score",
    )
    fill_box_cell(
        t.cell(0, 2),
        ["phoneme errors", "stress / syllable", "minimal-pair tip"],
        title="C  Coach output",
    )

    a = doc.add_table(rows=1, cols=1)
    arrow_cell(a.cell(0, 0), "↓  merge with today's STT score")

    t = doc.add_table(rows=1, cols=1)
    fill_box_cell(
        t.cell(0, 0),
        [
            "UI shows: wrong sound · unclear · skipped · stress",
            "Practice = listen TTS → speak → re-check with acoustic API",
            "Cost control: full passage stays on Whisper; only hard words hit Azure",
        ],
        title="LEARNER SEES",
        bg=LIGHT_BG,
    )
    para(doc, "", size=4, space_after=2)


def build() -> Path:
    out_dir = Path(__file__).resolve().parent
    out_path = out_dir / "Architecture.docx"
    alt_path = out_dir / "Architecture_v2.docx"

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Pronunciation Coach  ·  System Architecture  ·  July 2026")
        set_run_font(run, size=8, color=MUTED, italic=True)

    # --- title ---
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Pronunciation Coach")
    set_run_font(r, size=20, bold=True, color=ACCENT, name="Georgia")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("System architecture notes")
    set_run_font(r, size=12, italic=True, color=MUTED, name="Georgia")
    add_horizontal_rule(sub)

    para(
        doc,
        "How the pieces connect, why these APIs, how scoring works, DPDP handling, "
        "and what I'd build next for real pronunciation — not just STT text matching.",
        size=10,
        italic=True,
        color=MUTED,
        space_after=6,
    )

    # Live links up top — easy for evaluators
    para(doc, "Live links", size=10, bold=True, color=ACCENT, space_after=2)
    link_line(doc, "App (Vercel)", LIVE_APP)
    link_line(doc, "API (Render)", LIVE_API + "/health")
    link_line(doc, "Source", GITHUB)
    para(doc, "", size=4, space_after=2)

    # --- 1. Components ---
    heading(doc, "1. Components & how they connect")

    para(
        doc,
        "Two deployables. Frontend on Vercel, backend on Render. The browser hits "
        "the Vercel origin; Next.js rewrites /api/* to Render so CORS stays simple.",
        space_after=6,
    )

    add_current_architecture_diagram(doc)

    para(
        doc,
        "One attempt: consent → session → record/upload → analyze → score + highlights "
        "→ practice a few words → next round. Restart deletes the session.",
        space_after=4,
    )

    # --- 2. Models ---
    heading(doc, "2. Models & APIs (and why not the alternatives)")

    bullet(
        doc,
        "Groq whisper-large-v3 — fast, word timestamps + confidence. Skipped Azure "
        "Pronunciation Assessment for v1 (better phonemes, more setup/billing).",
    )
    bullet(
        doc,
        "g2p-en + RapidFuzz — local phoneme/string distance. No extra API. Deliberate shortcut.",
    )
    bullet(doc, "edge-tts — free TTS so the learner can hear the target word in practice.")
    bullet(
        doc,
        "OpenRouter (optional) — one short feedback line. Scoring works without it.",
    )
    bullet(
        doc,
        "Chunked STT on long audio — Whisper on 50–75s clips was truncating / inventing "
        "endings. Split into ~24s overlapping windows and merge.",
    )

    # --- 3. Scoring ---
    heading(doc, "3. How I score pronunciation & what I highlight")

    para(
        doc,
        "Not a lab-grade phoneme scorer. Goal: useful coach — show where speech drifted, give something to practice.",
        space_after=4,
    )

    para(doc, "Read Passage (scripted)", size=10.5, bold=True, space_after=2, color=ACCENT)
    para(
        doc,
        "Align expected words to STT. Per-word score mixes text similarity, phoneme "
        "similarity, and Whisper confidence. Missed words = 0. Highlight: wrong word, "
        "skip, low confidence, or trap words (island / comfortable) where STT hears a common mis-sounding.",
        space_after=4,
    )

    para(doc, "Free Speech & Jam", size=10.5, bold=True, space_after=2, color=ACCENT)
    para(
        doc,
        "No forced passage. Flag mumbled / low-confidence words; Jam also flags repeats "
        "and stuttery loops. Practice list capped so the learner isn't stuck.",
        space_after=4,
    )

    para(
        doc,
        "Honest limit: if someone mispronounces but Whisper still writes the correct "
        "spelling with high confidence, I may miss it. That's exactly what the next-week plan fixes.",
        size=10,
        italic=True,
        color=MUTED,
        space_after=4,
    )

    # --- 4. DPDP ---
    heading(doc, "4. DPDP Act 2023 — audio & personal data")

    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Consent", "Standalone notice before the test. Session only after agree. No account."),
        ("Storage", "Audio processed in memory for the request. Not written to disk / S3 / DB."),
        ("Retention", "In-memory session (~24h idle purge). Restart clears sooner."),
        ("Deletion", "DELETE /api/session + Restart in UI. Closing the tab drops client state."),
        ("Residency", "STT/TTS go to third-party APIs (US). Notice says this; purpose = assessment only."),
    ]
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.cell(i, 0), table.cell(i, 1)
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        set_run_font(r0, size=9.5, bold=True, color=ACCENT)
        r1 = c1.paragraphs[0].add_run(v)
        set_run_font(r1, size=9.5, color=INK)
        shade_cell(c0, LIGHT_BG)
        for cell in (c0, c1):
            set_cell_border(
                cell,
                top={"sz": "4", "color": "CCCCCC"},
                bottom={"sz": "4", "color": "CCCCCC"},
                left={"sz": "4", "color": "CCCCCC"},
                right={"sz": "4", "color": "CCCCCC"},
            )

    para(doc, "", size=4, space_after=2)
    para(
        doc,
        "No name/email/login. Voice is the sensitive bit — process, return the score, "
        "forget the bytes. Third-party STT is the main residency trade-off.",
        space_after=4,
    )

    # --- 5. Trade-offs + next week (expanded) ---
    heading(doc, "5. Trade-offs & what I'd build with one more week")

    para(doc, "Trade-offs I knowingly made", size=10.5, bold=True, space_after=2, color=ACCENT)
    bullet(doc, "STT + alignment over Azure phoneme scoring — shippable and cheap for v1.")
    bullet(doc, "In-memory sessions — fine for one Render instance; not multi-replica safe.")
    bullet(doc, "Render free cold starts — first hit can be slow.")
    bullet(
        doc,
        "Dropped trap words STT can't verify (whether/weather, castle/casual) — fair practice > clever traps.",
    )

    para(
        doc,
        "How I'd work on proper pronunciation (next week)",
        size=10.5,
        bold=True,
        space_before=8,
        space_after=4,
        color=ACCENT,
    )
    para(
        doc,
        "Today we mostly check “did Whisper hear the right word?” Real pronunciation "
        "is “did the mouth make the right sounds?” — even when the spelling comes out correct. "
        "Plan: keep Layer 1 (alignment) as the cheap pass, then add an acoustic pass only "
        "where it matters.",
        space_after=6,
    )

    add_next_week_diagram(doc)

    para(doc, "Day-by-day sketch", size=10, bold=True, space_after=2, color=ACCENT)
    bullet(
        doc,
        "Days 1–2 — Wire Azure Pronunciation Assessment (or Speechace) behind a feature "
        "flag. Cut each flagged word's audio using Whisper timestamps; send that clip only.",
    )
    bullet(
        doc,
        "Days 3–4 — Merge acoustic phoneme errors into our mistake types (wrong_sound, "
        "unclear, stress). Show tip like: “missing /θ/ in think — tongue between teeth.”",
    )
    bullet(
        doc,
        "Days 5–6 — Minimal-pair drills for the user's repeated errors (ship/sheep, "
        "think/sink). Practice verify uses the acoustic API, not STT spelling alone.",
    )
    bullet(
        doc,
        "Day 7 — Warm-up ping on landing (kill cold start), Redis sessions, optional "
        "India-region STT note in the DPDP notice if we add a regional endpoint.",
    )

    para(
        doc,
        "Why not acoustic on the whole passage from day one? Cost and latency. "
        "Whisper already gives alignment for free; Azure on every word would burn "
        "quota and slow the demo. Flagged words only is the production-shaped path.",
        size=10,
        italic=True,
        color=MUTED,
        space_after=6,
    )

    para(
        doc,
        "— end of notes —",
        size=9,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6,
        space_after=0,
    )

    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        # File open in Word/Cursor — write a sibling copy instead
        doc.save(alt_path)
        return alt_path


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
